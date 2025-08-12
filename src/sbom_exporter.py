import json
import csv
import logging
import uuid
from typing import Dict
from datetime import datetime

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

def get_value(value, default="unavailable"):
    """Returns the value if not None or empty, otherwise returns the default."""
    if value is None or (isinstance(value, str) and not value.strip()):
        return default
    return value


class SBOMExporter:
    @staticmethod
    def export_json(sbom, filename):
        # Process the SBOM to remove 'cast:' prefixes from property names
        processed_sbom = SBOMExporter._process_sbom_properties(sbom)
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(processed_sbom, f, indent=2)
        logging.info(f"SBOM exported to {filename}")
        
    @staticmethod
    def _process_sbom_properties(sbom):
        """Remove 'cast:' prefix from property names in the SBOM"""
        # Create a deep copy to avoid modifying the original
        processed_sbom = json.loads(json.dumps(sbom))
        
        # Process each component
        for component in processed_sbom.get("components", []):
            # Process properties
            if "properties" in component:
                new_properties = []
                for prop in component["properties"]:
                    if "name" in prop and get_value(prop["name"]).startswith("cast:"):
                        prop["name"] = get_value(prop["name"]).replace("cast:", "")
                    new_properties.append(prop)
                component["properties"] = new_properties
                        
        return processed_sbom

    @staticmethod
    def export_cyclonedx(sbom_data: Dict, filename: str, format: str = "json"):
        """Export SBOM data in CycloneDX format (manual implementation)"""
        try:
            # Create CycloneDX structure manually
            cyclonedx_bom = {
                "bomFormat": "CycloneDX",
                "specVersion": "1.4",
                "version": 1,
                "metadata": {
                    "timestamp": get_value(sbom_data.get("metadata", {}).get("timestamp"), datetime.utcnow().isoformat()),
                    "tools": [
                        {
                            "vendor": "CAST Highlight SBOM Generator",
                            "name": "SBOM Generator",
                            "version": get_value(sbom_data.get("metadata", {}).get("version"), "1.0")
                        }
                    ]
                },
                "components": []
            }

            # Add application as metadata component
            app_info = sbom_data.get("metadata", {}).get("application", {})
            if app_info:
                cyclonedx_bom["metadata"]["component"] = {
                    "type": "application",
                    "name": get_value(app_info.get("name")),
                    "version": get_value(app_info.get("version")),
                    "description": get_value(app_info.get("description"))
                }

            # Convert components to CycloneDX format
            for comp_data in sbom_data.get("components", []):
                cyclonedx_component = {
                    "type": "library",
                    "name": get_value(comp_data.get("name")),
                    "version": get_value(comp_data.get("version")),
                    "description": get_value(comp_data.get("description"))
                }
                
                # Add copyright if available
                if get_value(comp_data.get("copyright")) != "unavailable":
                    cyclonedx_component["copyright"] = get_value(comp_data.get("copyright"))

                # Add PURL
                if get_value(comp_data.get("purl")) != "unavailable":
                    cyclonedx_component["purl"] = get_value(comp_data["purl"])

                # Add licenses
                licenses = []
                for license_info in comp_data.get("licenses", []):
                    license_obj = {
                        "id": get_value(license_info.get("licenseId")),
                        "name": get_value(license_info.get("name"))
                    }
                    if get_value(license_info.get("url")) != "unavailable":
                        license_obj["url"] = get_value(license_info["url"])
                    licenses.append(license_obj)

                if licenses:
                    cyclonedx_component["licenses"] = licenses

                # Add external references
                ext_refs = []
                for ref in comp_data.get("externalReferences", []):
                    ref_type = "other"
                    if get_value(ref.get("type")) == "repository":
                        ref_type = "vcs"
                    elif get_value(ref.get("type")) == "website":
                        ref_type = "website"

                    ext_ref = {
                        "type": ref_type,
                        "url": get_value(ref.get("url"))
                    }
                    ext_refs.append(ext_ref)

                if ext_refs:
                    cyclonedx_component["externalReferences"] = ext_refs

                # Add properties
                properties = []
                for prop in comp_data.get("properties", []):
                    # Skip properties that are already represented in other fields
                    if get_value(prop.get("name")) in ["cast:timestamp"]:
                        continue
                        
                    # Ensure checksum/hash properties are properly formatted
                    if get_value(prop.get("name")) in ["cast:checksum", "cast:hash", "cast:fingerprint", "cast:sha1", "cast:sha256", "cast:md5"]:
                        hash_type = get_value(prop.get("name")).replace("cast:", "")
                        if "hashes" not in cyclonedx_component:
                            cyclonedx_component["hashes"] = []
                        cyclonedx_component["hashes"].append({
                            "alg": hash_type.upper(),
                            "content": get_value(prop.get("value"))
                        })
                    else:
                        # Remove 'cast:' prefix from property names
                        prop_name = get_value(prop.get("name"))
                        if prop_name.startswith("cast:"):
                            prop_name = prop_name.replace("cast:", "")
                            
                        property_obj = {
                            "name": prop_name,
                            "value": get_value(prop.get("value"))
                        }
                        properties.append(property_obj)

                if properties:
                    cyclonedx_component["properties"] = properties

                # Add vulnerabilities
                vulnerabilities = []
                for vuln_data in comp_data.get("vulnerabilities", []):
                    vuln = {
                        "id": get_value(vuln_data.get("id")),
                        "description": get_value(vuln_data.get("description"))
                    }

                    # Add CWE if available
                    if get_value(vuln_data.get("cweId")) != "unavailable":
                        vuln["cwes"] = [get_value(vuln_data["cweId"])]

                    # Add CPE if available
                    if get_value(vuln_data.get("cpe")) != "unavailable":
                        vuln["cpe"] = get_value(vuln_data["cpe"])

                    # Add rating if CVSS score is available
                    if get_value(vuln_data.get("cvssScore")) != "unavailable":
                        try:
                            cvss_score = float(get_value(vuln_data["cvssScore"]))
                            severity = "low"
                            if cvss_score >= 9.0:
                                severity = "critical"
                            elif cvss_score >= 7.0:
                                severity = "high"
                            elif cvss_score >= 4.0:
                                severity = "medium"

                            vuln["ratings"] = [
                                {
                                    "source": {
                                        "name": "CVSS",
                                        "url": "https://www.first.org/cvss/"
                                    },
                                    "score": cvss_score,
                                    "severity": severity,
                                    "method": "CVSSv3"
                                }
                            ]
                        except (ValueError, TypeError):
                            pass

                    # Add reference if available
                    if get_value(vuln_data.get("link")) != "unavailable":
                        vuln["references"] = [
                            {
                                "id": get_value(vuln_data.get("id")),
                                "source": {
                                    "name": "CVE",
                                    "url": get_value(vuln_data["link"])
                                }
                            }
                        ]

                    vulnerabilities.append(vuln)

                if vulnerabilities:
                    cyclonedx_component["vulnerabilities"] = vulnerabilities

                cyclonedx_bom["components"].append(cyclonedx_component)

            # Export based on format
            if format.lower() == "json":
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(cyclonedx_bom, f, indent=2, ensure_ascii=False)

                logging.info(f"CycloneDX JSON SBOM exported to {filename}")

            elif format.lower() == "xml":
                # Generate XML format
                xml_content = SBOMExporter._cyclonedx_to_xml(cyclonedx_bom)
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(xml_content)

                logging.info(f"CycloneDX XML SBOM exported to {filename}")

            else:
                logging.error(f"Unsupported CycloneDX format: {format}")
                return

        except Exception as e:
            logging.error(f"Failed to export CycloneDX: {e}")
            # Fallback to CLI method if manual method fails
            logging.info("Falling back to CLI method...")
            SBOMExporter.export_cyclonedx_with_cli(filename.replace(f".{format}", ""), format=format)

    @staticmethod
    def _cyclonedx_to_xml(cyclonedx_bom):
        """Convert CycloneDX BOM to XML format"""
        xml_lines = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<bom xmlns="http://cyclonedx.org/schema/bom/1.4"',
            '     version="1"',
            '     serialNumber="urn:uuid:' + str(uuid.uuid4()) + '">'
        ]

        # Add metadata
        if "metadata" in cyclonedx_bom:
            xml_lines.append('  <metadata>')
            
            if "timestamp" in cyclonedx_bom["metadata"]:
                xml_lines.append(f'    <timestamp>{get_value(cyclonedx_bom["metadata"]["timestamp"])}</timestamp>')
            
            if "tools" in cyclonedx_bom["metadata"]:
                xml_lines.append('    <tools>')
                for tool in cyclonedx_bom["metadata"]["tools"]:
                    xml_lines.append('      <tool>')
                    if "vendor" in tool:
                        xml_lines.append(f'        <vendor>{get_value(tool["vendor"])}</vendor>')
                    if "name" in tool:
                        xml_lines.append(f'        <name>{get_value(tool["name"])}</name>')
                    if "version" in tool:
                        xml_lines.append(f'        <version>{get_value(tool["version"])}</version>')
                    xml_lines.append('      </tool>')
                xml_lines.append('    </tools>')
            
            if "component" in cyclonedx_bom["metadata"]:
                comp = cyclonedx_bom["metadata"]["component"]
                xml_lines.append('    <component type="' + get_value(comp.get("type"), "application") + '">')
                if "name" in comp:
                    xml_lines.append(f'      <name>{get_value(comp["name"])}</name>')
                if "version" in comp:
                    xml_lines.append(f'      <version>{get_value(comp["version"])}</version>')
                if "description" in comp:
                    xml_lines.append(f'      <description>{get_value(comp["description"])}</description>')
                xml_lines.append('    </component>')
            
            xml_lines.append('  </metadata>')

        # Add components
        if "components" in cyclonedx_bom:
            xml_lines.append('  <components>')
            for comp in cyclonedx_bom["components"]:
                xml_lines.append('    <component type="' + get_value(comp.get("type"), "library") + '">')
                
                if "name" in comp:
                    xml_lines.append(f'      <name>{get_value(comp["name"])}</name>')
                if "version" in comp:
                    xml_lines.append(f'      <version>{get_value(comp["version"])}</version>')
                if "description" in comp:
                    xml_lines.append(f'      <description>{get_value(comp["description"])}</description>')
                if "purl" in comp:
                    xml_lines.append(f'      <purl>{get_value(comp["purl"])}</purl>')
                
                # Add licenses
                if "licenses" in comp:
                    xml_lines.append('      <licenses>')
                    for license_info in comp["licenses"]:
                        xml_lines.append('        <license>')
                        if "id" in license_info:
                            xml_lines.append(f'          <id>{get_value(license_info["id"])}</id>')
                        if "name" in license_info:
                            xml_lines.append(f'          <name>{get_value(license_info["name"])}</name>')
                        if "url" in license_info:
                            xml_lines.append(f'          <url>{get_value(license_info["url"])}</url>')
                        xml_lines.append('        </license>')
                    xml_lines.append('      </licenses>')
                
                # Add external references
                if "externalReferences" in comp:
                    xml_lines.append('      <externalReferences>')
                    for ref in comp["externalReferences"]:
                        xml_lines.append('        <reference>')
                        if "type" in ref:
                            xml_lines.append(f'          <type>{get_value(ref["type"])}</type>')
                        if "url" in ref:
                            xml_lines.append(f'          <url>{get_value(ref["url"])}</url>')
                        xml_lines.append('        </reference>')
                    xml_lines.append('      </externalReferences>')
                
                # Add properties
                if "properties" in comp:
                    xml_lines.append('      <properties>')
                    for prop in comp["properties"]:
                        xml_lines.append('        <property>')
                        if "name" in prop:
                            xml_lines.append(f'          <name>{get_value(prop["name"])}</name>')
                        if "value" in prop:
                            xml_lines.append(f'          <value>{get_value(prop["value"])}</value>')
                        xml_lines.append('        </property>')
                    xml_lines.append('      </properties>')
                
                # Add vulnerabilities
                if "vulnerabilities" in comp:
                    xml_lines.append('      <vulnerabilities>')
                    for vuln in comp["vulnerabilities"]:
                        xml_lines.append('        <vulnerability>')
                        if "id" in vuln:
                            xml_lines.append(f'          <id>{get_value(vuln["id"])}</id>')
                        if "description" in vuln:
                            xml_lines.append(f'          <description>{get_value(vuln["description"])}</description>')
                        
                        # Add CWEs
                        if "cwes" in vuln:
                            xml_lines.append('          <cwes>')
                            for cwe in vuln["cwes"]:
                                xml_lines.append(f'            <cwe>{get_value(cwe)}</cwe>')
                            xml_lines.append('          </cwes>')
                        
                        # Add ratings
                        if "ratings" in vuln:
                            xml_lines.append('          <ratings>')
                            for rating in vuln["ratings"]:
                                xml_lines.append('            <rating>')
                                if "score" in rating:
                                    xml_lines.append(f'              <score>{get_value(rating["score"])}</score>')
                                if "severity" in rating:
                                    xml_lines.append(f'              <severity>{get_value(rating["severity"])}</severity>')
                                if "method" in rating:
                                    xml_lines.append(f'              <method>{get_value(rating["method"])}</method>')
                                xml_lines.append('            </rating>')
                            xml_lines.append('          </ratings>')
                        
                        xml_lines.append('        </vulnerability>')
                    xml_lines.append('      </vulnerabilities>')
                
                xml_lines.append('    </component>')
            xml_lines.append('  </components>')

        xml_lines.append('</bom>')
        return '\n'.join(xml_lines)

    @staticmethod
    def export_cyclonedx_with_cli(output_path: str, format: str = "json"):
        """Export CycloneDX using CLI tool (fallback method)"""
        try:
            import subprocess
            import sys
            
            # Try to use cyclonedx-py CLI tool
            cmd = [sys.executable, "-m", "cyclonedx_py", "bom", "--input-format", "json", "--output-format", format, "--output-file", f"{output_path}.{format}"]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                logging.info(f"CycloneDX {format.upper()} exported using CLI to {output_path}.{format}")
            else:
                logging.error(f"CLI export failed: {result.stderr}")
                
        except Exception as e:
            logging.error(f"CLI export error: {e}")

    @staticmethod
    def export_docx(sbom_data: Dict, filename: str):
        """Export SBOM as DOCX document"""
        if not Document:
            logging.error('python-docx is not installed. Cannot export to .docx.')
            return
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Software Bill of Materials (SBOM)', 0)
            title.alignment = 1  # Center alignment
            
            # Add metadata section
            doc.add_heading('Metadata', level=1)
            metadata = sbom_data.get("metadata", {})
            
            metadata_table = doc.add_table(rows=1, cols=2)
            metadata_table.style = 'Table Grid'
            hdr_cells = metadata_table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'
            
            # Add metadata rows
            metadata_fields = [
                ('Timestamp', get_value(metadata.get('timestamp'))),
                ('Tool', get_value(metadata.get('tool'), 'CAST Highlight SBOM Generator')),
                ('Version', get_value(metadata.get('version'), '2.0')),
                ('Total Components', str(len(sbom_data.get('components', []))))
            ]
            
            for field, value in metadata_fields:
                row_cells = metadata_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Add components section
            doc.add_heading('Components', level=1)
            
            if sbom_data.get("components"):
                # Create components table
                components_table = doc.add_table(rows=1, cols=6)
                components_table.style = 'Table Grid'
                hdr_cells = components_table.rows[0].cells
                hdr_cells[0].text = 'Name'
                hdr_cells[1].text = 'Version'
                hdr_cells[2].text = 'Type'
                hdr_cells[3].text = 'Description'
                hdr_cells[4].text = 'License'
                hdr_cells[5].text = 'Vulnerabilities'
                
                # Add component rows
                for component in sbom_data["components"]:
                    row_cells = components_table.add_row().cells
                    row_cells[0].text = get_value(component.get('name'))
                    row_cells[1].text = get_value(component.get('version'))
                    row_cells[2].text = get_value(component.get('type'))
                    description = get_value(component.get('description'))
                    row_cells[3].text = description[:100] + '...' if len(description) > 100 else description
                    
                    # License information
                    licenses = component.get('licenses', [])
                    license_text = '; '.join([get_value(lic.get('name'), '') for lic in licenses]) if licenses else 'unavailable'
                    row_cells[4].text = license_text
                    
                    # Vulnerability count
                    vuln_count = len(component.get('vulnerabilities', []))
                    row_cells[5].text = str(vuln_count)
            
            # Add summary section
            doc.add_heading('Summary', level=1)
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"This SBOM contains {len(sbom_data.get('components', []))} components. ")
            summary_para.add_run("Generated by CAST Highlight SBOM Generator v2.0.")
            
            # Save document
            doc.save(filename)
            logging.info(f"SBOM exported to {filename}")
            
        except Exception as e:
            logging.error(f"Failed to export DOCX: {e}")

    @staticmethod
    def export_csv(sbom_data: Dict, filename: str):
        """Export SBOM as CSV with all fields"""
        processed_sbom = SBOMExporter._process_sbom_properties(sbom_data)
        with open(filename, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Write header
            writer.writerow([
                "Component Name", "Component Version", "Component Type", "Component Description", "PURL",
                "Copyright", "License", "External References", "Hashes",
                "Vulnerability ID", "Vulnerability Description", "Vulnerability CWE", "Vulnerability CPE",
                "Vulnerability CVSS Score", "Vulnerability Severity", "Vulnerability Link"
            ])
            for component in processed_sbom["components"]:
                hashes = "; ".join([f"{get_value(h['alg'])}:{get_value(h['content'])}" for h in component.get("hashes", [])])
                for i, vuln in enumerate(component.get("vulnerabilities", [])):
                    row = [
                        get_value(component.get("name")),
                        get_value(component.get("version")),
                        get_value(component.get("type")),
                        get_value(component.get("description")),
                        get_value(component.get("purl")),
                        get_value(component.get("copyright")),
                        "; ".join([get_value(lic.get("name")) for lic in component.get("licenses", [])]),
                        "; ".join([get_value(ref.get("url")) for ref in component.get("externalReferences", [])]),
                        hashes,
                        get_value(vuln.get("id")),
                        get_value(vuln.get("description")),
                        "; ".join([get_value(cwe) for cwe in vuln.get("cwes", [])]),
                        get_value(vuln.get("cpe")),
                        get_value(vuln.get("ratings", [{}])[0].get("score")),
                        get_value(vuln.get("ratings", [{}])[0].get("severity")),
                        get_value(vuln.get("references", [{}])[0].get("url"))
                    ]
                    if i > 0:
                        # for subsequent vulnerabilities of the same component, clear component data
                        row[:9] = [""] * 9
                    writer.writerow(row)
                if not component.get("vulnerabilities"):
                    writer.writerow([
                        get_value(component.get("name")),
                        get_value(component.get("version")),
                        get_value(component.get("type")),
                        get_value(component.get("description")),
                        get_value(component.get("purl")),
                        get_value(component.get("copyright")),
                        "; ".join([get_value(lic.get("name")) for lic in component.get("licenses", [])]),
                        "; ".join([get_value(ref.get("url")) for ref in component.get("externalReferences", [])]),
                        hashes,
                        "", "", "", "", "", "", ""
                    ])
        logging.info(f"SBOM exported to {filename}")


    @staticmethod
    def export_xlsx(sbom_data: Dict, filename: str):
        """Export SBOM as XLSX with all fields"""
        if not openpyxl:
            logging.error('openpyxl is not installed. Cannot export to .xlsx.')
            return
        
        processed_sbom = SBOMExporter._process_sbom_properties(sbom_data)
        wb = openpyxl.Workbook()
        
        # Components sheet
        ws = wb.active
        ws.title = "SBOM Components"
        
        # Collect all property names
        all_properties = set()
        for component in processed_sbom["components"]:
            for prop in component.get("properties", []):
                all_properties.add(get_value(prop["name"]))
        
        header = [
            "Component Name", "Version", "Type", "Description", "PURL",
            "Copyright", "License", "External References", "Hashes"
        ] + sorted(list(all_properties))
        ws.append(header)

        for component in processed_sbom["components"]:
            hashes = "; ".join([f"{get_value(h['alg'])}:{get_value(h['content'])}" for h in component.get("hashes", [])])
            
            properties_map = {get_value(prop["name"]): get_value(prop["value"]) for prop in component.get("properties", [])}
            
            row = [
                get_value(component.get("name")),
                get_value(component.get("version")),
                get_value(component.get("type")),
                get_value(component.get("description")),
                get_value(component.get("purl")),
                get_value(component.get("copyright")),
                "; ".join([get_value(lic.get("name")) for lic in component.get("licenses", [])]),
                "; ".join([get_value(ref.get("url")) for ref in component.get("externalReferences", [])]),
                hashes
            ] + [properties_map.get(prop_name, "unavailable") for prop_name in sorted(list(all_properties))]
            ws.append(row)

        # Vulnerabilities sheet
        ws_vuln = wb.create_sheet(title="Vulnerabilities")
        ws_vuln.append([
            "Component Name", "Vulnerability ID (CVE)", "Description", "CWE ID", "CPE", 
            "Rating Source", "CVSS Score", "Severity", "Rating Method", "Reference ID", "Reference URL"
        ])
        for component in processed_sbom["components"]:
            for vuln in component.get("vulnerabilities", []):
                for rating in vuln.get("ratings", [{}]):
                    for ref in vuln.get("references", [{}]):
                        ws_vuln.append([
                            get_value(component.get("name")),
                            get_value(vuln.get("id")),
                            get_value(vuln.get("description")),
                            "; ".join([get_value(cwe) for cwe in vuln.get("cwes", [])]),
                            get_value(vuln.get("cpe")),
                            get_value(rating.get("source", {}).get("name")),
                            get_value(rating.get("score")),
                            get_value(rating.get("severity")),
                            get_value(rating.get("method")),
                            get_value(ref.get("id")),
                            get_value(ref.get("source", {}).get("url"))
                        ])
        
        wb.save(filename)
        logging.info(f"SBOM exported to {filename}")