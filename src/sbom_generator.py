#!/usr/bin/env python3
"""
SBOM Generator for CAST Highlight Compliance
This application extracts software component data from CAST Highlight API
and generates SBOM (Software Bill of Materials) compliant with industry standards.
"""

import requests
import json
import csv
import sys
from datetime import datetime
from typing import Dict, List, Optional
import logging
import os
import subprocess

# For new output formats
try:
    import openpyxl  # For Excel
    from openpyxl import styles
except ImportError:
    openpyxl = None
    styles = None
try:
    from docx import Document  # For DOCX
except ImportError:
    Document = None

# Import logging configuration
from logging_config import setup_module_logging

# Set up separated logging for sbom_generator module
logger, log_files = setup_module_logging('sbom_generator')

# CycloneDX support (JSON/XML)
try:
    from cyclonedx.model.bom import Bom
    from cyclonedx.model.component import Component, ComponentType
except ImportError as e:
    Bom = None
    Component = None
    ComponentType = None
    logger.error(
        f"cyclonedx-python-lib is not installed or import failed: {e}. Cannot export CycloneDX."
    )


def load_config(config_path: str) -> dict:
    if not os.path.exists(config_path):
        logger.error(f"Config file {config_path} not found.")
        sys.exit(1)
    with open(config_path, "r") as f:
        config = json.load(f)
    return config


class HighlightAPI:
    """Client for CAST Highlight API"""

    def __init__(
        self,
        base_url: str,
        username: str | None = None,
        password: str | None = None,
        api_key: str | None = None,
        company_id: str | None = None,
    ):
        # Accept base_url with or without /WS2, but do not append /WS2 if already present
        base_url = base_url.rstrip("/")
        if not base_url.endswith("/WS2"):
            base_url = base_url + "/WS2"
        self.base_url = base_url
        self.username = username
        self.password = password
        self.api_key = api_key
        self.company_id = company_id
        self.access_token = None

        # Initialize headers
        self.headers = {"Content-Type": "application/json"}

        # Authenticate if credentials provided
        if username and password:
            self._authenticate_with_credentials()
        elif api_key:
            self._authenticate_with_api_key()

    def _authenticate_with_credentials(self):
        """Authenticate using username/password with Basic HTTP Authentication"""
        try:
            # Use Basic HTTP Authentication instead of token-based auth
            from requests.auth import HTTPBasicAuth

            if self.username and self.password:
                self.auth = HTTPBasicAuth(self.username, self.password)
                logger.info(
                    "Successfully set up Basic HTTP Authentication with username/password"
                )
            else:
                raise Exception(
                    "Username and password are required for Basic HTTP Authentication"
                )

        except Exception as e:
            logger.error(f"Authentication setup failed: {e}")
            raise Exception(f"Authentication setup failed: {e}")

    def _authenticate_with_api_key(self):
        """Authenticate using API key"""
        self.headers["Authorization"] = f"Bearer {self.api_key}"
        logger.info("Using API key authentication")

    def get_applications(self) -> List[Dict]:
        """Get list of applications from CAST Highlight"""
        try:
            url = f"{self.base_url}/domains/{self.company_id}/applications/"
            # Use auth parameter for Basic HTTP Authentication
            response = requests.get(
                url, headers=self.headers, auth=getattr(self, "auth", None)
            )
            response.raise_for_status()
            data = response.json()
            # Ensure we return a list
            if isinstance(data, list):
                return data
            elif isinstance(data, dict) and "applications" in data:
                return data["applications"]
            else:
                logger.warning(f"Unexpected applications data structure: {type(data)}")
                return []
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to get applications: {e}")
            return []

    def get_application_details(self, app_id: str) -> Optional[Dict]:
        """Get detailed information about a specific application"""
        try:
            url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}"
            response = requests.get(
                url, headers=self.headers, auth=getattr(self, "auth", None)
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to get application details for {app_id}: {e}")
            return None

    def get_components(self, app_id: str) -> List[Dict]:
        """Get components/dependencies for an application"""
        # Try the dependencies endpoint first (based on API documentation)
        try:
            url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/dependencies"
            response = requests.get(
                url, headers=self.headers, auth=getattr(self, "auth", None)
            )
            response.raise_for_status()
            data = response.json()

            # Handle different response structures
            if isinstance(data, dict) and "dependencies" in data:
                return data["dependencies"]
            elif isinstance(data, dict) and "thirdParties" in data:
                return data["thirdParties"]
            elif isinstance(data, list):
                return data
            else:
                logger.warning(f"Unexpected component data structure for {app_id}")
                return []

        except requests.exceptions.RequestException as e:
            logger.warning(f"Dependencies endpoint failed for {app_id}: {e}")

            # Fallback to thirdparty endpoint
            try:
                url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/thirdparty"
                response = requests.get(
                    url, headers=self.headers, auth=getattr(self, "auth", None)
                )
                response.raise_for_status()
                data = response.json()
                if isinstance(data, dict) and "thirdParties" in data:
                    return data["thirdParties"]
                elif isinstance(data, list):
                    return data
                elif isinstance(data, dict):
                    return [data]
                return []
            except requests.exceptions.RequestException as e2:
                logger.error(
                    f"Both dependencies and thirdparty endpoints failed for {app_id}: {e2}"
                )
                return []

    def get_vulnerabilities(self, app_id: str) -> List[Dict]:
        """Get security vulnerabilities for an application"""
        # Use the aggregated vulnerabilities endpoint
        try:
            url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/vulnerabilities/aggregated/"
            response = requests.get(
                url, headers=self.headers, auth=getattr(self, "auth", None)
            )
            response.raise_for_status()
            data = response.json()

            # If the response contains vulnerabilities in a nested structure, extract them
            if isinstance(data, dict) and "vulnerabilities" in data:
                return data["vulnerabilities"]
            elif isinstance(data, list):
                return data
            else:
                logger.warning(f"Unexpected vulnerability data structure for {app_id}")
                return []

        except requests.exceptions.RequestException as e:
            logger.warning(
                f"Aggregated vulnerabilities endpoint failed for {app_id}: {e}"
            )

            # Fallback to regular vulnerabilities endpoint
            try:
                url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/vulnerabilities"
                response = requests.get(
                    url, headers=self.headers, auth=getattr(self, "auth", None)
                )
                response.raise_for_status()
                return response.json()
            except requests.exceptions.RequestException as e2:
                logger.warning(
                    f"Regular vulnerabilities endpoint also failed for {app_id}: {e2}"
                )

                # Final fallback to CVE endpoint
                try:
                    url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/cve"
                    response = requests.get(
                        url, headers=self.headers, auth=getattr(self, "auth", None)
                    )
                    response.raise_for_status()
                    return response.json()
                except requests.exceptions.RequestException as e3:
                    logger.error(
                        f"All vulnerability endpoints failed for {app_id}: {e3}"
                    )
                    return []

    def get_licenses(self, app_id: str) -> List[Dict]:
        """Get license information for an application"""
        try:
            url = f"{self.base_url}/domains/{self.company_id}/applications/{app_id}/licenses"
            response = requests.get(
                url, headers=self.headers, auth=getattr(self, "auth", None)
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to get licenses for {app_id}: {e}")
            return []


def get_value(value, default="unavailable"):
    """Returns the value if not None or empty, otherwise returns the default."""
    if value is None or (isinstance(value, str) and not value.strip()):
        return default
    return value


class SBOMGenerator:
    """Generate SBOM data from CAST Highlight information"""

    def __init__(self, cast_api: HighlightAPI):
        self.cast_api = cast_api
        self.sbom_data = {
            "sbomVersion": "1.0",
            "metadata": {
                "timestamp": datetime.utcnow().isoformat(),
                "tool": "CAST Highlight SBOM Generator",
                "version": "1.0",
            },
            "components": [],
        }

    def generate_sbom(self, app_id: str) -> Dict:
        """Generate complete SBOM for an application"""
        logger.info(f"Generating SBOM for application {app_id}")

        # Get application details
        app_details = self.cast_api.get_application_details(app_id)
        if app_details:
            self.sbom_data["metadata"]["application"] = {
                "name": get_value(app_details.get("name")),
                "version": get_value(app_details.get("version")),
                "description": get_value(app_details.get("description")),
            }

        # Get components
        components = self.cast_api.get_components(app_id)
        for component in components:
            sbom_component = self._convert_to_sbom_component(component)
            if sbom_component:
                self.sbom_data["components"].append(sbom_component)

        # Get vulnerabilities and add to components
        vulnerabilities = self.cast_api.get_vulnerabilities(app_id)
        self._add_vulnerabilities_to_components(vulnerabilities)

        # Get licenses and add to components
        licenses = self.cast_api.get_licenses(app_id)
        self._add_licenses_to_components(licenses)

        logger.info(
            f"Generated SBOM with {len(self.sbom_data['components'])} components"
        )
        return self.sbom_data

    def _convert_to_sbom_component(self, cast_component: Dict) -> Optional[Dict]:
        """Convert CAST Highlight component to SBOM format"""
        try:
            # Extract vulnerabilities from component's cve field
            vulnerabilities = []
            if cast_component.get("cve") and cast_component["cve"].get(
                "vulnerabilities"
            ):
                for vuln in cast_component["cve"]["vulnerabilities"]:
                    vulnerabilities.append(
                        {
                            "id": get_value(vuln.get("name")),
                            "description": get_value(vuln.get("description")),
                            "severity": get_value(vuln.get("criticity")),
                            "cvssScore": get_value(vuln.get("cvssScore")),
                            "cweId": get_value(vuln.get("cweId")),
                            "cpe": get_value(vuln.get("cpe")),
                            "isKev": vuln.get("isKev", False),
                            "link": get_value(vuln.get("link")),
                        }
                    )

            return {
                "type": "library",
                "name": get_value(cast_component.get("name")),
                "version": get_value(cast_component.get("version")),
                "description": get_value(cast_component.get("description")),
                "purl": self._generate_purl(cast_component),
                "externalReferences": self._get_external_references(cast_component),
                "properties": self._get_component_properties(cast_component),
                "supplier": {
                    "name": get_value(cast_component.get("supplierName")),
                    "contact": get_value(cast_component.get("supplierContact")),
                },
                "author": get_value(cast_component.get("author")),
                "copyright": get_value(cast_component.get("copyright")),
                "licenses": [],  # Will be populated later
                "vulnerabilities": vulnerabilities,  # Extract from component's cve field
            }
        except Exception as e:
            logger.error(
                f"Error converting component {cast_component.get('name', 'unavailable')}: {e}"
            )
            return None

    def _generate_purl(self, component: Dict) -> str:
        """Generate Package URL (PURL) for component"""
        name = get_value(component.get("name"))
        version = get_value(component.get("version"))
        package_type = get_value(component.get("packageType"), default="generic")

        if package_type == "maven":
            return f"pkg:maven/{name}@{version}"
        elif package_type == "npm":
            return f"pkg:npm/{name}@{version}"
        elif package_type == "nuget":
            return f"pkg:nuget/{name}@{version}"
        elif package_type == "pypi":
            return f"pkg:pypi/{name}@{version}"
        else:
            return f"pkg:generic/{name}@{version}"

    def _get_external_references(self, component: Dict) -> List[Dict]:
        """Get external references for component"""
        references = []

        if get_value(component.get("repositoryUrl")) != "unavailable":
            references.append({"type": "repository", "url": get_value(component["repositoryUrl"])})

        if get_value(component.get("homepageUrl")) != "unavailable":
            references.append({"type": "website", "url": get_value(component["homepageUrl"])})

        return references

    def _get_component_properties(self, component: Dict) -> List[Dict]:
        """Get component properties"""
        properties = []

        # Add CAST Highlight specific properties
        if get_value(component.get("packageType")) != "unavailable":
            properties.append(
                {"name": "cast:packageType", "value": get_value(component["packageType"])}
            )

        if get_value(component.get("filePath")) != "unavailable":
            properties.append({"name": "cast:filePath", "value": get_value(component["filePath"])}
            )

        if get_value(component.get("origin")) != "unavailable":
            properties.append({"name": "cast:origin", "value": get_value(component["origin"])}
            )

        if get_value(component.get("dependencies")) != "unavailable":
            properties.append({"name": "cast:dependencies", "value": get_value(component["dependencies"])}
            )

        if get_value(component.get("patchStatus")) != "unavailable":
            properties.append({"name": "cast:patchStatus", "value": get_value(component["patchStatus"])}
            )

        if get_value(component.get("releaseDate")) != "unavailable":
            properties.append({"name": "cast:releaseDate", "value": get_value(component["releaseDate"])}
            )

        if get_value(component.get("eolDate")) != "unavailable":
            properties.append({"name": "cast:eolDate", "value": get_value(component["eolDate"])}
            )

        if get_value(component.get("criticality")) != "unavailable":
            properties.append({"name": "cast:criticality", "value": get_value(component["criticality"])}
            )

        if get_value(component.get("usageRestrictions")) != "unavailable":
            properties.append({"name": "cast:usageRestrictions", "value": get_value(component["usageRestrictions"])}
            )

        if get_value(component.get("checksum")) != "unavailable":
            properties.append({"name": "cast:checksum", "value": get_value(component["checksum"])}
            )

        if get_value(component.get("comments")) != "unavailable":
            properties.append({"name": "cast:comments", "value": get_value(component["comments"])}
            )

        if get_value(component.get("executable")) != "unavailable":
            properties.append({"name": "cast:executable", "value": get_value(component["executable"])}
            )

        if get_value(component.get("archive")) != "unavailable":
            properties.append({"name": "cast:archive", "value": get_value(component["archive"])}
            )

        if get_value(component.get("structured")) != "unavailable":
            properties.append({"name": "cast:structured", "value": get_value(component["structured"])}
            )

        return properties

    def _add_vulnerabilities_to_components(self, vulnerabilities: List[Dict]):
        """Add vulnerability information to components"""
        for vuln in vulnerabilities:
            component_name = get_value(vuln.get("componentName"))
            if component_name != "unavailable":
                # Find matching component
                for component in self.sbom_data["components"]:
                    if component["name"] == component_name:
                        component["vulnerabilities"].append(
                            {
                                "id": get_value(vuln.get("cveId")),
                                "description": get_value(vuln.get("description")),
                                "severity": get_value(vuln.get("severity")),
                                "cvssScore": get_value(vuln.get("cvssScore")),
                                "publishedDate": get_value(vuln.get("publishedDate")),
                                "references": vuln.get("references", []),
                            }
                        )
                        break

    def _add_licenses_to_components(self, licenses: List[Dict]):
        """Add license information to components"""
        for license_info in licenses:
            component_name = get_value(license_info.get("componentName"))
            if component_name != "unavailable":
                # Find matching component
                for component in self.sbom_data["components"]:
                    if component["name"] == component_name:
                        component["licenses"].append(
                            {
                                "licenseId": get_value(license_info.get("licenseId")),
                                "name": get_value(license_info.get("licenseName")),
                                "url": get_value(license_info.get("licenseUrl")),
                                "compliance": get_value(license_info.get("compliance")),
                            }
                        )
                        break


class SBOMExporter:
    """Export SBOM data in various formats"""

    @staticmethod
    def _process_sbom_properties(sbom_data: Dict) -> Dict:
        """Process SBOM data to remove cast: prefixes from property names"""
        import copy
        processed_data = copy.deepcopy(sbom_data)
        
        for component in processed_data.get("components", []):
            # Process properties to remove cast: prefix
            if "properties" in component:
                for prop in component["properties"]:
                    if "name" in prop and prop["name"].startswith("cast:"):
                        prop["name"] = prop["name"][5:]  # Remove "cast:" prefix
        
        return processed_data

    @staticmethod
    def export_json(sbom_data: Dict, filename: str):
        # Process SBOM data to remove cast: prefixes
        processed_data = SBOMExporter._process_sbom_properties(sbom_data)
        
        # Ensure all fields present in Excel export are included in JSON
        for component in processed_data["components"]:
            properties = {
                prop.get("name", ""): prop.get("value", "")
                for prop in component.get("properties", [])
            }
            # Add all Excel fields to the component dict if not present
            component["supplier_name"] = component.get("supplier", {}).get("name", "")
            component["component_license"] = "; ".join(
                [lic.get("name", "") for lic in component.get("licenses", [])]
            )
            component["component_origin"] = properties.get("origin", "unavailable")
            component["component_dependencies"] = properties.get(
                "dependencies", "unavailable"
            )
            component["vulnerabilities_count"] = str(
                len(component.get("vulnerabilities", []))
            )
            component["patch_status"] = properties.get("patchStatus", "unavailable")
            component["release_date"] = properties.get("releaseDate", "")
            component["eol_date"] = properties.get("eolDate", "")
            component["criticality"] = properties.get("criticality", "unavailable")
            component["usage_restrictions"] = properties.get(
                "usageRestrictions", "None"
            )
            component["checksums"] = properties.get("checksum", "")
            component["comments"] = properties.get("comments", "")
            component["author_of_sbom_data"] = component.get("author", "")
            component["timestamp"] = processed_data.get("metadata", {}).get("timestamp", "")
            component["executable_property"] = properties.get("executable", "No")
            component["archive_property"] = properties.get("archive", "No")
            component["structured_property"] = properties.get("structured", "No")
            component["external_references"] = "; ".join(
                [ref.get("url", "") for ref in component.get("externalReferences", [])]
            )
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(processed_data, f, indent=2)
        logger.info(f"SBOM exported to {filename}")

    @staticmethod
    def export_csv(sbom_data: Dict, filename: str):
        """Export SBOM as CSV with all baseline fields (parity with Excel export)"""
        # Process SBOM data to remove cast: prefixes
        processed_data = SBOMExporter._process_sbom_properties(sbom_data)
        
        with open(filename, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            # Write header (same as Components Complete in Excel)
            writer.writerow(
                [
                    "Component Name",
                    "Component Version",
                    "Component Description",
                    "Component Supplier",
                    "Component License",
                    "Component Origin",
                    "Component Dependencies",
                    "Vulnerabilities Count",
                    "Patch Status",
                    "Release Date",
                    "End-of-Life (EOL) Date",
                    "Criticality",
                    "Usage Restrictions",
                    "Checksums or Hashes",
                    "Comments or Notes",
                    "Author of SBOM Data",
                    "Timestamp",
                    "Executable Property",
                    "Archive Property",
                    "Structured Property",
                    "Unique Identifier (PURL)",
                    "Component Type",
                    "Copyright",
                    "External References",
                ]
            )
            for component in processed_data["components"]:
                properties = {
                    prop.get("name", ""): prop.get("value", "")
                    for prop in component.get("properties", [])
                }
                writer.writerow(
                    [
                        component.get("name", ""),
                        component.get("version", ""),
                        component.get("description", ""),
                        component.get("supplier", {}).get("name", ""),
                        "; ".join(
                            [
                                lic.get("name", "")
                                for lic in component.get("licenses", [])
                            ]
                        ),
                        properties.get("origin", "unavailable"),
                        properties.get("dependencies", "unavailable"),
                        str(len(component.get("vulnerabilities", []))),
                        properties.get("patchStatus", "unavailable"),
                        properties.get("releaseDate", ""),
                        properties.get("eolDate", ""),
                        properties.get("criticality", "unavailable"),
                        properties.get("usageRestrictions", "None"),
                        properties.get("checksum", ""),
                        properties.get("comments", ""),
                        component.get("author", ""),
                        processed_data.get("metadata", {}).get("timestamp", ""),
                        properties.get("executable", "No"),
                        properties.get("archive", "No"),
                        properties.get("structured", "No"),
                        component.get("purl", ""),
                        component.get("type", ""),
                        component.get("copyright", ""),
                        "; ".join(
                            [
                                ref.get("url", "")
                                for ref in component.get("externalReferences", [])
                            ]
                        ),
                    ]
                )
        logger.info(f"SBOM exported to {filename}")

    @staticmethod
    def export_spdx(sbom_data: Dict, filename: str):
        """Export SBOM in SPDX format with license ID filtering and robust error handling"""
        try:
            def is_valid_spdx_license_id(license_id):
                # SPDX license IDs are usually at least 3 characters and alphanumeric (e.g., MIT, Apache-2.0)
                # This can be improved with a full SPDX list if needed
                return isinstance(license_id, str) and len(license_id) > 2 and all(c.isalnum() or c in ['-', '.'] for c in license_id)

            spdx_content = f"""SPDXVersion: SPDX-2.2
DataLicense: CC0-1.0
SPDXID: SPDXRef-DOCUMENT
DocumentName: {sbom_data.get('metadata', {}).get('application', {}).get('name', 'unavailable')} SBOM
DocumentNamespace: http://spdx.org/spdxdocs/cast-highlight-sbom-{datetime.now().strftime('%Y%m%d-%H%M%S')}
Creator: Tool: CAST Highlight SBOM Generator
Created: {datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')}

"""

            for i, component in enumerate(sbom_data["components"], 1):
                # Filter license IDs
                license_ids = [lic.get('licenseId', 'NOASSERTION') for lic in component.get('licenses', [])]
                filtered_license_ids = [lid for lid in license_ids if is_valid_spdx_license_id(lid)]
                if not filtered_license_ids:
                    filtered_license_ids = ['NOASSERTION']

                spdx_content += f"""PackageName: {component.get('name', 'unavailable')}
SPDXID: SPDXRef-Package-{i}
PackageVersion: {component.get('version', 'unavailable')}
PackageDownloadLocation: NOASSERTION
PackageLicenseConcluded: {'; '.join(filtered_license_ids)}
PackageLicenseDeclared: {'; '.join(filtered_license_ids)}
PackageCopyrightText: {component.get('copyright', 'NOASSERTION')}
PackageDescription: {component.get('description', 'NOASSERTION')}

"""

            with open(filename, "w", encoding="utf-8") as f:
                f.write(spdx_content)

            logger.info(f"SBOM exported to {filename}")
        except Exception as e:
            logger.error(f"Failed to export SPDX: {e}")
            try:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(f"# SPDX export failed: {e}\n")
            except Exception as inner_e:
                logger.error(f"Failed to write error message to SPDX file: {inner_e}")

    @staticmethod
    def export_xlsx(sbom_data: Dict, filename: str):
        if not openpyxl:
            logger.error("openpyxl is not installed. Cannot export to .xlsx.")
            return

        # Process SBOM data to remove cast: prefixes
        processed_data = SBOMExporter._process_sbom_properties(sbom_data)

        wb = openpyxl.Workbook()

        # Remove default sheet and create our own
        if wb.active:
            wb.remove(wb.active)

        # 1. SBOM Overview/Summary Sheet
        ws_summary = wb.create_sheet(title="SBOM Summary")
        ws_summary.append(["SBOM Information", "Value"])
        ws_summary.append(
            ["SBOM Version", processed_data.get("sbomVersion", "1.0")]
        )
        ws_summary.append(
            ["Generated Timestamp", processed_data.get("metadata", {}).get("timestamp", "")]
        )
        ws_summary.append(["Tool", processed_data.get("metadata", {}).get("tool", "")])
        ws_summary.append(
            ["Tool Version", processed_data.get("metadata", {}).get("version", "")]
        )
        ws_summary.append(
            [
                "Application Name",
                processed_data.get("metadata", {}).get("application", {}).get("name", ""),
            ]
        )
        ws_summary.append(
            [
                "Application Version",
                processed_data.get("metadata", {}).get("application", {}).get("version", ""),
            ]
        )
        ws_summary.append(["Total Components", len(processed_data.get("components", []))])

        # Count vulnerabilities by severity
        vuln_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
        for component in processed_data.get("components", []):
            for vuln in component.get("vulnerabilities", []):
                severity = vuln.get("severity", "UNKNOWN")
                if severity in vuln_counts:
                    vuln_counts[severity] += 1

        ws_summary.append(["Critical Vulnerabilities", vuln_counts["CRITICAL"]])
        ws_summary.append(["High Vulnerabilities", vuln_counts["HIGH"]])
        ws_summary.append(["Medium Vulnerabilities", vuln_counts["MEDIUM"]])
        ws_summary.append(["Low Vulnerabilities", vuln_counts["LOW"]])

        # 2. Complete Components Sheet (All baseline information)
        ws_components = wb.create_sheet(title="Components Complete")
        ws_components.append(
            [
                "Component Name",
                "Component Version",
                "Component Description",
                "Component Supplier",
                "Component License",
                "Component Origin",
                "Component Dependencies",
                "Vulnerabilities Count",
                "Patch Status",
                "Release Date",
                "End-of-Life (EOL) Date",
                "Criticality",
                "Usage Restrictions",
                "Checksums or Hashes",
                "Comments or Notes",
                "Author of SBOM Data",
                "Timestamp",
                "Executable Property",
                "Archive Property",
                "Structured Property",
                "Unique Identifier (PURL)",
                "Component Type",
                "Copyright",
                "External References",
            ]
        )

        for component in processed_data.get("components", []):
            # Extract properties for additional fields
            properties = {
                prop.get("name", ""): prop.get("value", "")
                for prop in component.get("properties", [])
            }

            ws_components.append(
                [
                    component.get("name", ""),
                    component.get("version", ""),
                    component.get("description", ""),
                    component.get("supplier", {}).get("name", ""),
                    "; ".join(
                        [lic.get("name", "") for lic in component.get("licenses", [])]
                    ),
                    properties.get("origin", "unavailable"),  # Component Origin
                    properties.get(
                        "dependencies", "unavailable"
                    ),  # Component Dependencies
                    str(len(component.get("vulnerabilities", []))),
                    properties.get("patchStatus", "unavailable"),  # Patch Status
                    properties.get("releaseDate", ""),  # Release Date
                    properties.get("eolDate", ""),  # EOL Date
                    properties.get("criticality", "unavailable"),  # Criticality
                    properties.get(
                        "usageRestrictions", "None"
                    ),  # Usage Restrictions
                    properties.get("checksum", ""),  # Checksums
                    properties.get("comments", ""),  # Comments
                    component.get("author", ""),
                    processed_data.get("metadata", {}).get("timestamp", ""),
                    properties.get("executable", "No"),  # Executable Property
                    properties.get("archive", "No"),  # Archive Property
                    properties.get("structured", "No"),  # Structured Property
                    component.get("purl", ""),
                    component.get("type", ""),
                    component.get("copyright", ""),
                    "; ".join(
                        [
                            ref.get("url", "")
                            for ref in component.get("externalReferences", [])
                        ]
                    ),
                ]
            )

        # 3. Vulnerabilities Detailed Sheet
        ws_vuln = wb.create_sheet(title="Vulnerabilities")
        ws_vuln.append(
            [
                "Component Name",
                "Component Version",
                "Vulnerability ID (CVE)",
                "Severity",
                "CVSS Score",
                "Description",
                "CWE ID",
                "CPE",
                "Reference/Link",
                "Is KEV",
                "Published Date",
                "Patch Status",
                "Remediation Notes",
            ]
        )

        for component in processed_data.get("components", []):
            for vuln in component.get("vulnerabilities", []):
                ws_vuln.append(
                    [
                        component.get("name", ""),
                        component.get("version", ""),
                        vuln.get("id", ""),
                        vuln.get("severity", ""),
                        vuln.get("cvssScore", ""),
                        vuln.get("description", ""),
                        vuln.get("cweId", ""),
                        vuln.get("cpe", ""),
                        vuln.get("link", ""),
                        "Yes" if vuln.get("isKev", False) else "No",
                        vuln.get("publishedDate", ""),
                        vuln.get("patchStatus", "unavailable"),
                        vuln.get("remediationNotes", ""),
                    ]
                )

        # 4. Licenses Sheet
        ws_licenses = wb.create_sheet(title="Licenses")
        ws_licenses.append(
            [
                "Component Name",
                "Component Version",
                "License Name",
                "License Compliance",
                "License Type",
                "License URL",
                "License Text",
                "SPDX Identifier",
            ]
        )

        for component in processed_data.get("components", []):
            for license_info in component.get("licenses", []):
                ws_licenses.append(
                    [
                        component.get("name", ""),
                        component.get("version", ""),
                        license_info.get("name", ""),
                        license_info.get("compliance", "unavailable"),
                        license_info.get("type", ""),
                        license_info.get("url", ""),
                        license_info.get("text", ""),
                        license_info.get("spdxId", ""),
                    ]
                )

        # 5. Dependencies Sheet
        ws_deps = wb.create_sheet(title="Dependencies")
        ws_deps.append(
            [
                "Component Name",
                "Component Version",
                "Dependency Name",
                "Dependency Version",
                "Dependency Type",
                "Relationship Type",
                "Scope",
                "Optional",
            ]
        )

        for component in processed_data.get("components", []):
            # Extract dependency information from properties
            properties = {
                prop.get("name", ""): prop.get("value", "")
                for prop in component.get("properties", [])
            }
            dependencies = properties.get("dependencies", "")

            if dependencies and dependencies != "unavailable":
                # Parse dependencies if available
                deps_list = dependencies.split(";")
                for dep in deps_list:
                    ws_deps.append(
                        [
                            component.get("name", ""),
                            component.get("version", ""),
                            dep.strip(),
                            "",
                            "library",
                            "depends_on",
                            "runtime",
                            "No",
                        ]
                    )
            else:
                # Add a placeholder row
                ws_deps.append(
                    [
                        component.get("name", ""),
                        component.get("version", ""),
                        "No explicit dependencies found",
                        "",
                        "",
                        "",
                        "",
                        "",
                    ]
                )

        # 6. Security Analysis Sheet
        ws_security = wb.create_sheet(title="Security Analysis")
        ws_security.append(
            [
                "Component Name",
                "Component Version",
                "Total Vulnerabilities",
                "Critical Count",
                "High Count",
                "Medium Count",
                "Low Count",
                "KEV Count",
                "Risk Score",
                "Recommendation",
                "Last Security Review",
            ]
        )

        for component in sbom_data.get("components", []):
            vulns = component.get("vulnerabilities", [])
            critical_count = sum(1 for v in vulns if v.get("severity") == "CRITICAL")
            high_count = sum(1 for v in vulns if v.get("severity") == "HIGH")
            medium_count = sum(1 for v in vulns if v.get("severity") == "MEDIUM")
            low_count = sum(1 for v in vulns if v.get("severity") == "LOW")
            kev_count = sum(1 for v in vulns if v.get("isKev", False))

            # Calculate risk score (simple formula)
            risk_score = (
                (critical_count * 10)
                + (high_count * 7)
                + (medium_count * 4)
                + (low_count * 1)
            )

            # Generate recommendation
            if critical_count > 0:
                recommendation = (
                    "IMMEDIATE ACTION REQUIRED - Critical vulnerabilities found"
                )
            elif high_count > 0:
                recommendation = (
                    "HIGH PRIORITY - High severity vulnerabilities need attention"
                )
            elif medium_count > 0:
                recommendation = (
                    "MEDIUM PRIORITY - Consider updating or replacing component"
                )
            elif low_count > 0:
                recommendation = "LOW PRIORITY - Monitor for updates"
            else:
                recommendation = "No known vulnerabilities"

            ws_security.append(
                [
                    component.get("name", ""),
                    component.get("version", ""),
                    len(vulns),
                    critical_count,
                    high_count,
                    medium_count,
                    low_count,
                    kev_count,
                    risk_score,
                    recommendation,
                    sbom_data.get("metadata", {}).get("timestamp", ""),
                ]
            )

        # 7. Metadata Sheet
        ws_metadata = wb.create_sheet(title="Metadata")
        ws_metadata.append(["Field", "Value", "Source", "Last Updated"])

        metadata = sbom_data.get("metadata", {})
        ws_metadata.append(
            [
                "SBOM Version",
                sbom_data.get("sbomVersion", ""),
                "Generated",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Tool",
                metadata.get("tool", ""),
                "Generated",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Tool Version",
                metadata.get("version", ""),
                "Generated",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Application Name",
                metadata.get("application", {}).get("name", ""),
                "CAST Highlight",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Application Version",
                metadata.get("application", {}).get("version", ""),
                "CAST Highlight",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Application Description",
                metadata.get("application", {}).get("description", ""),
                "CAST Highlight",
                metadata.get("timestamp", ""),
            ]
        )
        ws_metadata.append(
            [
                "Total Components",
                str(len(sbom_data.get("components", []))),
                "Generated",
                metadata.get("timestamp", ""),
            ]
        )

        # Format worksheets for better readability
        from openpyxl.cell.cell import MergedCell
        for ws in wb.worksheets:
            # Auto-fit columns
            for col in ws.columns:
                max_length = 0
                column_letter = None
                for cell in col:
                    if not isinstance(cell, MergedCell):
                        if column_letter is None:
                            column_letter = cell.column_letter
                        try:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        except Exception:
                            pass
                if column_letter:
                    adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
                    ws.column_dimensions[column_letter].width = adjusted_width

            # Bold headers
            if styles:
                for cell in ws[1]:
                    cell.font = styles.Font(bold=True)
                    cell.fill = styles.PatternFill(
                        start_color="CCCCCC",
                        end_color="CCCCCC",
                        fill_type="solid"
                    )

        wb.save(filename)
        logger.info(f"SBOM exported to {filename} with {len(wb.worksheets)} worksheets")

    @staticmethod
    def export_cyclonedx(sbom_data: Dict, filename: str, format: str = "json"):
        """Export SBOM data in CycloneDX format (manual implementation)"""
        try:
            # Process SBOM data to remove cast: prefixes
            processed_data = SBOMExporter._process_sbom_properties(sbom_data)
            
            # Create CycloneDX structure manually
            cyclonedx_bom = {
                "bomFormat": "CycloneDX",
                "specVersion": "1.4",
                "version": 1,
                "metadata": {
                    "timestamp": processed_data.get("metadata", {}).get("timestamp", datetime.utcnow().isoformat()),
                    "tools": [
                        {
                            "vendor": "CAST Highlight SBOM Generator",
                            "name": "SBOM Generator",
                            "version": processed_data.get("metadata", {}).get("version", "1.0")
                        }
                    ]
                },
                "components": []
            }
            
            # Add application as metadata component
            app_info = processed_data.get("metadata", {}).get("application", {})
            if app_info:
                cyclonedx_bom["metadata"]["component"] = {
                    "type": "application",
                    "name": app_info.get("name", "unavailable"),
                    "version": app_info.get("version", "unavailable"),
                    "description": app_info.get("description", "unavailable")
                }
            
            # Convert components to CycloneDX format
            for comp_data in processed_data.get("components", []):
                cyclonedx_component = {
                    "type": "library",
                    "name": comp_data.get("name", "unavailable"),
                    "version": comp_data.get("version", "unavailable"),
                    "description": comp_data.get("description", "unavailable")
                }
                
                # Add PURL
                if comp_data.get("purl"):
                    cyclonedx_component["purl"] = comp_data["purl"]
                
                # Add licenses
                licenses = []
                for license_info in comp_data.get("licenses", []):
                    license_obj = {
                        "id": license_info.get("licenseId", "unavailable"),
                        "name": license_info.get("name", "unavailable")
                    }
                    if license_info.get("url"):
                        license_obj["url"] = license_info["url"]
                    licenses.append(license_obj)
                
                if licenses:
                    cyclonedx_component["licenses"] = licenses
                
                # Add external references
                ext_refs = []
                for ref in comp_data.get("externalReferences", []):
                    ref_type = "other"
                    if ref.get("type") == "repository":
                        ref_type = "vcs"
                    elif ref.get("type") == "website":
                        ref_type = "website"
                    
                    ext_ref = {
                        "type": ref_type,
                        "url": ref.get("url", "unavailable")
                    }
                    ext_refs.append(ext_ref)
                
                if ext_refs:
                    cyclonedx_component["externalReferences"] = ext_refs
                
                # Add properties
                properties = []
                for prop in comp_data.get("properties", []):
                    property_obj = {
                        "name": prop.get("name", "unavailable"),
                        "value": prop.get("value", "unavailable")
                    }
                    properties.append(property_obj)
                
                if properties:
                    cyclonedx_component["properties"] = properties
                
                # Add vulnerabilities
                vulnerabilities = []
                for vuln_data in comp_data.get("vulnerabilities", []):
                    vuln = {
                        "id": vuln_data.get("id", "unavailable"),
                        "description": vuln.get("description", "unavailable")
                    }
                    
                    # Add CWE if available
                    if vuln_data.get("cweId"):
                        vuln["cwes"] = [vuln_data["cweId"]]
                    
                    # Add CPE if available
                    if vuln_data.get("cpe"):
                        vuln["cpe"] = vuln_data["cpe"]
                    
                    # Add rating if CVSS score is available
                    if vuln_data.get("cvssScore"):
                        try:
                            cvss_score = float(vuln_data["cvssScore"])
                            severity = "low"
                            if cvss_score >= 9.0:
                                severity = "critical"
                            elif cvss_score >= 7.0:
                                severity = "high"
                            elif cvss_score >= 4.0:
                                severity = "medium"
                            
                            vuln["ratings"] = [
                                {
                                    "source": {
                                        "name": "CVSS",
                                        "url": "https://www.first.org/cvss/"
                                    },
                                    "score": cvss_score,
                                    "severity": severity,
                                    "method": "CVSSv3"
                                }
                            ]
                        except (ValueError, TypeError):
                            pass
                    
                    # Add reference if available
                    if vuln_data.get("link"):
                        vuln["references"] = [
                            {
                                "id": vuln_data.get("id", "unavailable"),
                                "source": {
                                    "name": "CVE",
                                    "url": vuln_data["link"]
                                }
                            }
                        ]
                    
                    vulnerabilities.append(vuln)
                
                if vulnerabilities:
                    cyclonedx_component["vulnerabilities"] = vulnerabilities
                
                cyclonedx_bom["components"].append(cyclonedx_component)
            
            # Export based on format
            if format.lower() == "json":
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(cyclonedx_bom, f, indent=2, ensure_ascii=False)
                
                logger.info(f"CycloneDX JSON SBOM exported to {filename}")
                
            elif format.lower() == "xml":
                # Generate XML format
                xml_content = SBOMExporter._cyclonedx_to_xml(cyclonedx_bom)
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(xml_content)
                
                logger.info(f"CycloneDX XML SBOM exported to {filename}")
                
            else:
                logger.error(f"Unsupported CycloneDX format: {format}")
                return
                
        except Exception as e:
            logging.error(f"Failed to export CycloneDX: {e}")
            # Fallback to CLI method if manual method fails
            logger.info("Falling back to CLI method...")
            SBOMExporter.export_cyclonedx_with_cli(filename.replace(f".{format}", ""), format=format)

    @staticmethod
    def _cyclonedx_to_xml(cyclonedx_bom):
        """Convert CycloneDX BOM to XML format"""
        xml_lines = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<bom xmlns="http://cyclonedx.org/schema/bom/1.4"',
            '     version="1"',
            '     serialNumber="urn:uuid:' + SBOMExporter._generate_uuid() + '">'
        ]
        
        # Metadata
        metadata = cyclonedx_bom.get("metadata", {})
        xml_lines.append('  <metadata>')
        
        if metadata.get("timestamp"):
            xml_lines.append(f'    <timestamp>{metadata["timestamp"]}</timestamp>')
        
        # Tools
        if metadata.get("tools"):
            xml_lines.append('    <tools>')
            for tool in metadata["tools"]:
                xml_lines.append('      <tool>')
                xml_lines.append(f'        <vendor>{tool.get("vendor", "")}</vendor>')
                xml_lines.append(f'        <name>{tool.get("name", "")}</name>')
                xml_lines.append(f'        <version>{tool.get("version", "")}</version>')
                xml_lines.append('      </tool>')
            xml_lines.append('    </tools>')
        
        # Component
        if metadata.get("component"):
            comp = metadata["component"]
            xml_lines.append('    <component type="' + comp.get("type", "application") + '">')
            xml_lines.append(f'      <name>{comp.get("name", "")}</name>')
            xml_lines.append(f'      <version>{comp.get("version", "")}</version>')
            if comp.get("description"):
                xml_lines.append(f'      <description>{comp.get("description")}</description>')
            xml_lines.append('    </component>')
        
        xml_lines.append('  </metadata>')
        
        # Components
        if cyclonedx_bom.get("components"):
            xml_lines.append('  <components>')
            for comp in cyclonedx_bom["components"]:
                xml_lines.append('    <component type="' + comp.get("type", "library") + '">')
                xml_lines.append(f'      <name>{comp.get("name", "")}</name>')
                xml_lines.append(f'      <version>{comp.get("version", "")}</version>')
                if comp.get("description"):
                    xml_lines.append(f'      <description>{comp.get("description")}</description>')
                if comp.get("purl"):
                    xml_lines.append(f'      <purl>{comp.get("purl")}</purl>')
                
                # Licenses
                if comp.get("licenses"):
                    xml_lines.append('      <licenses>')
                    for license_info in comp["licenses"]:
                        xml_lines.append('        <license>')
                        xml_lines.append(f'          <id>{license_info.get("id", "")}</id>')
                        xml_lines.append(f'          <name>{license_info.get("name", "")}</name>')
                        if license_info.get("url"):
                            xml_lines.append(f'          <url>{license_info.get("url")}</url>')
                        xml_lines.append('        </license>')
                    xml_lines.append('      </licenses>')
                
                # External References
                if comp.get("externalReferences"):
                    xml_lines.append('      <externalReferences>')
                    for ref in comp["externalReferences"]:
                        xml_lines.append('        <reference type="' + ref.get("type", "other") + '">')
                        xml_lines.append(f'          <url>{ref.get("url", "")}</url>')
                        xml_lines.append('        </reference>')
                    xml_lines.append('      </externalReferences>')
                
                # Properties
                if comp.get("properties"):
                    xml_lines.append('      <properties>')
                    for prop in comp["properties"]:
                        xml_lines.append('        <property>')
                        xml_lines.append(f'          <name>{prop.get("name", "")}</name>')
                        xml_lines.append(f'          <value>{prop.get("value", "")}</value>')
                        xml_lines.append('        </property>')
                    xml_lines.append('      </properties>')
                
                # Vulnerabilities
                if comp.get("vulnerabilities"):
                    xml_lines.append('      <vulnerabilities>')
                    for vuln in comp["vulnerabilities"]:
                        xml_lines.append('        <vulnerability>')
                        xml_lines.append(f'          <id>{vuln.get("id", "")}</id>')
                        if vuln.get("description"):
                            xml_lines.append(f'          <description>{vuln.get("description")}</description>')
                        if vuln.get("cwes"):
                            xml_lines.append('          <cwes>')
                            for cwe in vuln["cwes"]:
                                xml_lines.append(f'            <cwe>{cwe}</cwe>')
                            xml_lines.append('          </cwes>')
                        if vuln.get("cpe"):
                            xml_lines.append(f'          <cpe>{vuln.get("cpe")}</cpe>')
                        xml_lines.append('        </vulnerability>')
                    xml_lines.append('      </vulnerabilities>')
                
                xml_lines.append('    </component>')
            xml_lines.append('  </components>')
        
        xml_lines.append('</bom>')
        
        return '\n'.join(xml_lines)

    @staticmethod
    def _generate_uuid():
        """Generate a simple UUID for CycloneDX"""
        import uuid
        return str(uuid.uuid4())

    @staticmethod
    def export_docx(sbom_data: Dict, filename: str):
        if not Document:
            logger.error("python-docx is not installed. Cannot export to .docx.")
            return
        
        # Process SBOM data to remove cast: prefixes
        processed_data = SBOMExporter._process_sbom_properties(sbom_data)
        
        doc = Document()
        doc.add_heading("Software Bill of Materials (SBOM)", 0)
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}")
        doc.add_heading("Components", level=1)
        # Match columns to Components Complete in Excel
        headers = [
            "Component Name",
            "Component Version",
            "Component Description",
            "Component Supplier",
            "Component License",
            "Component Origin",
            "Component Dependencies",
            "Vulnerabilities Count",
            "Patch Status",
            "Release Date",
            "End-of-Life (EOL) Date",
            "Criticality",
            "Usage Restrictions",
            "Checksums or Hashes",
            "Comments or Notes",
            "Author of SBOM Data",
            "Timestamp",
            "Executable Property",
            "Archive Property",
            "Structured Property",
            "Unique Identifier (PURL)",
            "Component Type",
            "Copyright",
            "External References",
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for component in processed_data["components"]:
            properties = {
                prop.get("name", ""): prop.get("value", "")
                for prop in component.get("properties", [])
            }
            row_cells = table.add_row().cells
            values = [
                component.get("name", ""),
                component.get("version", ""),
                component.get("description", ""),
                component.get("supplier", {}).get("name", ""),
                "; ".join(
                    [lic.get("name", "") for lic in component.get("licenses", [])]
                ),
                properties.get("origin", "unavailable"),
                properties.get("dependencies", "unavailable"),
                str(len(component.get("vulnerabilities", []))),
                properties.get("patchStatus", "unavailable"),
                properties.get("releaseDate", ""),
                properties.get("eolDate", ""),
                properties.get("criticality", "unavailable"),
                properties.get("usageRestrictions", "None"),
                properties.get("checksum", ""),
                properties.get("comments", ""),
                component.get("author", ""),
                processed_data.get("metadata", {}).get("timestamp", ""),
                properties.get("executable", "No"),
                properties.get("archive", "No"),
                properties.get("structured", "No"),
                component.get("purl", ""),
                component.get("type", ""),
                component.get("copyright", ""),
                "; ".join(
                    [
                        ref.get("url", "")
                        for ref in component.get("externalReferences", [])
                    ]
                ),
            ]
            for i, v in enumerate(values):
                row_cells[i].text = str(v)
        doc.save(filename)
        logger.info(f"SBOM exported to {filename}")

    @staticmethod
    def export_cyclonedx_with_cli(output_path: str, format: str = "json"):
        """Fallback method using cyclonedx-py CLI tool"""
        fmt = "JSON" if format == "json" else "XML"
        result = subprocess.run(
            [
                "cyclonedx-py",
                "environment",
                "--output-format",
                fmt,
                "--output-file",
                f"{output_path}_cyclonedx.{format}",
            ],
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            logger.error(f"CycloneDX CLI failed: {result.stderr}")
        else:
            logger.info(f"CycloneDX SBOM exported to {output_path}_cyclonedx.{format}")


def get_conf(conf, *keys, default=None):
    for k in keys:
        if k in conf:
            return conf[k]
    return default


def validate_application_id(cast_api: HighlightAPI, app_id: str) -> bool:
    """Return True if app_id exists in the list of applications, else False."""
    apps = cast_api.get_applications()
    if not apps:
        logger.error(
            "No applications found for the current user/company. Cannot validate Application ID."
        )
        return False
    for app in apps:
        if str(app.get("id")) == str(app_id):
            logger.info(
                f"Validated Application ID {app_id}: {app.get('name', '[no name]')}"
            )
            return True
    logger.error(
        f"Application ID {app_id} not found in your CAST Highlight company."
    )
    return False


def main():
    # Always use config/config.json in the config directory
    config_path = "config/config.json"
    config = load_config(config_path)
    ch_conf = config.get("cast_highlight", {})
    auth_conf = ch_conf.get("authentication", {})
    sbom_conf = config.get("sbom_settings", {})

    # Application ID or Name
    app_id = config.get("application_id") or sbom_conf.get("application_id")
    app_name = config.get("application_name") or sbom_conf.get("application_name")

    # Output formats (list)
    output_formats = sbom_conf.get("output_formats")
    if not output_formats or not isinstance(output_formats, list):
        output_formats = ["json"]

    # Authentication
    method = auth_conf.get("method", "api_key")
    company_id = (
        auth_conf.get("company_id")
        or ch_conf.get("company_id")
        or config.get("company_id")
    )
    if not company_id:
        logger.error("company_id is required in config for all authentication methods.")
        sys.exit(1)
    if method == "credentials":
        cast_api = HighlightAPI(
            ch_conf["api_url"],
            username=auth_conf.get("username"),
            password=auth_conf.get("password"),
            company_id=company_id,
        )
    else:
        cast_api = HighlightAPI(
            ch_conf["api_url"], api_key=auth_conf.get("api_key"), company_id=company_id
        )

    # If app_id is not provided, but app_name is, look up the ID
    if not app_id and app_name:
        logger.info(f"Looking up application ID for name: {app_name}")
        apps = cast_api.get_applications()
        match = next(
            (app for app in apps if app.get("name", "").lower() == app_name.lower()),
            None,
        )
        if match:
            app_id = match["id"]
            logger.info(f'Found application "{app_name}" with ID: {app_id}')
        else:
            logger.warning(
                f'Application name "{app_name}" not found in your CAST Highlight company. Falling back to application_id from config if available.'
            )
            app_id = config.get("application_id") or sbom_conf.get("application_id")
            if app_id:
                logger.info(f"Using fallback application_id: {app_id}")
            else:
                logger.error("No valid application_id found in config to fall back to.")
                sys.exit(1)
    if not app_id:
        logger.error(
            "Application ID or name is required. Please provide it in the config file."
        )
        sys.exit(1)

    # Generate SBOM
    generator = SBOMGenerator(cast_api)
    sbom_data = generator.generate_sbom(app_id)

    # Create Reports directory if it doesn't exist
    reports_dir = "Reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
        print(f"Created Reports directory: {reports_dir}")

    # Get application details for filename
    app_name = (
        sbom_data.get("metadata", {}).get("application", {}).get("name", "unavailable")
    )
    # Clean app name for filename (remove special characters)
    safe_app_name = "".join(
        c for c in app_name if c.isalnum() or c in (" ", "-", "_")
    ).rstrip()
    safe_app_name = safe_app_name.replace(" ", "_")

    # Create filename with app name, ID, and timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename_base = f"{safe_app_name}_ID{app_id}_{timestamp}"

    # Full path for output files
    output_path = os.path.join(reports_dir, filename_base)

    # Export in requested format(s)
    for fmt in output_formats:
        fmt = fmt.lower()
        if fmt == "json":
            SBOMExporter.export_json(sbom_data, f"{output_path}.json")
        elif fmt == "csv":
            SBOMExporter.export_csv(sbom_data, f"{output_path}.csv")
        elif fmt == "xlsx":
            SBOMExporter.export_xlsx(sbom_data, f"{output_path}.xlsx")
        elif fmt == "cyclonedx":
            SBOMExporter.export_cyclonedx_with_cli(output_path, format="json")
            SBOMExporter.export_cyclonedx_with_cli(output_path, format="xml")
        elif fmt == "docx":
            SBOMExporter.export_docx(sbom_data, f"{output_path}.docx")
        else:
            logger.warning(f"Unknown output format: {fmt}. Skipping.")

    print(
        f"SBOM generation completed successfully! Components found: {len(sbom_data['components'])}"
    )
    print(f"Files saved in Reports directory with prefix: {filename_base}")


if __name__ == "__main__":
    main()

